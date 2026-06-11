"""构建 tcp/udp_packet_capture.docx，对齐任务书风格。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import time

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
GIT_URL = BASE_DIR


def safe_save(doc, path):
    """保存文档，Word 占用时自动回退到带时间戳的文件名"""
    try:
        doc.save(path)
    except PermissionError:
        ts = int(time.time())
        alt = path.replace(".docx", f"_{ts}.docx")
        doc.save(alt)
        print(f"  (原文件被占用，已保存到: {os.path.basename(alt)})")


def set_run(run, cn="宋体", en="Times New Roman", size=12):
    run.font.size = Pt(size)
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    for k, v in [("w:eastAsia", cn), ("w:ascii", en), ("w:hAnsi", en), ("w:cs", en)]:
        rf.set(qn(k), v)
    rpr.insert(0, rf)


def sec(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(f"{label} {text}")
    set_run(r, "黑体", "Times New Roman", 12)
    r.bold = True


def para(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run(r, "宋体", "Times New Roman", 12)
    r.bold = bold


def code(doc, text):
    """带语法高亮的代码块"""
    from pygments import lex
    from pygments.lexers import PythonLexer
    from pygments.token import Token

    # 颜色映射（Pygments Token -> RGB）
    COLOR_MAP = {
        Token.Keyword:             "0000FF",  # 蓝
        Token.Keyword.Namespace:   "0000FF",
        Token.Keyword.Type:        "0000FF",
        Token.Name.Builtin:        "008080",  # 青
        Token.Name.Function:       "795E26",  # 棕
        Token.Name.Class:          "267F99",  # 深蓝
        Token.Name.Decorator:      "808000",  # 橄榄
        Token.String:              "A31515",  # 红
        Token.String.Affix:        "A31515",
        Token.String.Backtick:     "A31515",
        Token.String.Char:         "A31515",
        Token.String.Delimiter:    "A31515",
        Token.String.Doc:          "008000",  # 绿（docstring）
        Token.String.Double:       "A31515",
        Token.String.Escape:       "A31515",
        Token.String.Heredoc:      "A31515",
        Token.String.Interpol:     "A31515",
        Token.String.Other:        "A31515",
        Token.String.Regex:        "A31515",
        Token.String.Single:       "A31515",
        Token.String.Symbol:       "A31515",
        Token.Number:              "098658",  # 深绿
        Token.Comment:             "008000",  # 绿
        Token.Comment.Single:      "008000",
        Token.Comment.Multiline:   "008000",
        Token.Operator:            "000000",  # 黑
        Token.Punctuation:         "000000",
        Token.Name:                "000000",
        Token.Text:                "000000",
        Token.Name.Builtin.Pseudo: "008080",
    }

    def get_color(token_type):
        """从 token 类型找最近的颜色"""
        t = token_type
        while t is not None:
            if t in COLOR_MAP:
                return COLOR_MAP[t]
            t = t.parent
        return "000000"

    tokens = list(lex(text, PythonLexer()))

    # 创建代码表格
    t = doc.add_table(1, 1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    c._element.get_or_add_tcPr().append(shd)

    tcPr = c._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "999999")
        tcBorders.append(border)
    tcPr.append(tcBorders)

    # 逐 token 添加带颜色的 run（换行时新增段落）
    p = c.paragraphs[0]
    for ttype, tvalue in tokens:
        if tvalue == '\n':
            # 换行：在当前段落末尾加空 run（保留行结构），然后新建段落
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(11)
            continue

        color = get_color(ttype)
        r = p.add_run(tvalue)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        rf = OxmlElement("w:rFonts")
        for k, v in [("w:eastAsia", "宋体"), ("w:ascii", "Consolas"), ("w:hAnsi", "Consolas")]:
            rf.set(qn(k), v)
        r._element.get_or_add_rPr().insert(0, rf)
        # 设置颜色
        r.font.color.rgb = None  # clear default
        rPr = r._element.get_or_add_rPr()
        c_elem = OxmlElement("w:color")
        c_elem.set(qn("w:val"), color)
        rPr.append(c_elem)

    # 设置首段格式
    first_p = c.paragraphs[0]
    first_p.paragraph_format.space_before = Pt(4)
    first_p.paragraph_format.line_spacing = Pt(11)
    # 末段底边距
    last_p = c.paragraphs[-1]
    last_p.paragraph_format.space_after = Pt(4)


def cap(doc, text, img_path=None):
    """添加图片标题 + 嵌入图片（如果 img_path 存在）"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = 1
    r = p.add_run(text)
    set_run(r, "宋体", "Times New Roman", 10.5)

    if img_path and os.path.exists(img_path):
        img_p = doc.add_paragraph()
        img_p.alignment = 1
        img_p.paragraph_format.space_after = Pt(12)
        r = img_p.add_run()
        r.add_picture(img_path, width=Cm(14))


def build_task1():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(12)
    s.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    p = doc.add_paragraph()
    p.alignment = 1
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("TCP Reverse — Wireshark 抓包说明文档")
    set_run(r, "黑体", "Times New Roman", 16)
    r.bold = True

    task1_dir = os.path.join(BASE_DIR, "Task1")

    sec(doc, "a.", "Wireshark 报文捕获截图")
    para(doc, "以下截图展示四种报文类型在 Wireshark 中的实际捕获情况。Display Filter: tcp.port == <端口号>")
    cap(doc, "图1 — 四种报文概览（Initialization / agree / reverseRequest / reverseAnswer）",
        os.path.join(task1_dir, "overview.png"))
    cap(doc, "图2 — Type=1 Initialization 报文详情（Type + N 块数）",
        os.path.join(task1_dir, "type1_init.png"))
    cap(doc, "图3 — Type=2 agree 报文详情（单字节 Type=2）",
        os.path.join(task1_dir, "type2_agree.png"))
    cap(doc, "图4 — Type=3 reverseRequest 报文详情（Type + Length + Data）",
        os.path.join(task1_dir, "type3_request.png"))
    cap(doc, "图5 — Type=4 reverseAnswer 报文详情（Type + Length + 反转后 Data）",
        os.path.join(task1_dir, "type4_answer.png"))

    sec(doc, "b.", "实现上的关键点和对应的代码解决方案")
    para(doc, "关键点1：随机分块算法。使用 chunk_seed 固定随机种子，保证分块结果可复现。验收时给定 (文件大小, Lmin, Lmax, seed) 可手推 N 及各块长度。剩余字节 <= Lmax 时直接作为最后一块。代码位于 reversetcpclient.py 中的 split_file() 函数。", True)
    code(doc, """def split_file(data: bytes, lmin: int, lmax: int, seed: int) -> list:
    rng = random.Random(seed)
    total = len(data)
    chunks = []
    pos = 0
    while pos < total:
        remaining = total - pos
        if remaining <= lmax:
            chunks.append(remaining)
            break
        size = rng.randint(lmin, lmax)
        chunks.append(size)
        pos += size
    return chunks""")
    para(doc, "关键点2：自定义协议报文设计与交互。使用 struct.pack/unpack 构造和解析 4 种类型报文。Type 字段 1 字节（!B），Length 字段 4 字节（!I），统一 Big-Endian。TCP 流式传输使用 recv_exact() 精确接收防止粘包。代码位于 common.py。", True)
    code(doc, """TYPE_INIT = 1     # Initialization: [1B Type][4B N]
TYPE_AGREE = 2    # agree:          [1B Type]
TYPE_REQUEST = 3  # reverseRequest: [1B Type][4B Length][Data]
TYPE_ANSWER = 4   # reverseAnswer:  [1B Type][4B Length][Data]

init = struct.pack("!BI", TYPE_INIT, n_blocks)
msg_type, n = struct.unpack("!BI", header)
conn.sendall(answer)
data = recv_exact(sock, n)""")
    para(doc, "关键点3：多客户端并发 + 线程安全日志。Server 使用 threading.Thread 为每个连接创建独立处理线程。run_log.txt 写入使用 threading.Lock 保证多线程写安全，每条日志带毫秒级时间戳，与 Wireshark 抓包时间可相互印证。", True)
    code(doc, """_LOG_LOCK = threading.Lock()

def log_event(log_path, fmt, *args):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
    line = f"[{timestamp}] {fmt.format(*args)}"
    with _LOG_LOCK:
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(line + "\\n")

while True:
    conn, addr = server_socket.accept()
    t = threading.Thread(target=handle_client, args=(conn, addr), daemon=True)
    t.start()""")
    para(doc, "关键点4：common.py + config.py 模块化设计。将 recv_exact、log_event、TYPE_* 常量提取到 common.py，server 和 client 共同 import，消除代码重复。同时将网络配置、协议常量、分块默认值统一管理在 config.py 中。", True)

    sec(doc, "c.", "掌握的知识点")
    knowledge = [
        "TCP Socket 编程：socket() -> bind() -> listen() -> accept() 服务端流程；"
        "socket() -> connect() 客户端流程；TCP 全双工、面向连接、可靠传输特性。",
        "应用层自定义报文：struct 模块进行 Big-Endian 序列化/反序列化；"
        "Type-Length-Data 报文结构设计原则；报文类型区分与状态机交互。",
        "TCP 流式传输与粘包：sendall() 保证全量发送；recv_exact() 循环接收解决粘包问题；"
        "理解 TCP 面向字节流与 UDP 面向报文的本质区别。",
        "多线程并发编程：threading.Thread 为每个连接创建线程；"
        "threading.Lock 互斥锁保护共享资源（日志文件）；daemon 线程自动回收。",
        "分块算法与随机种子：random.Random(seed) 确定性随机；"
        "Offset 累加分块定位；边界条件处理（最后块 <= Lmax）。",
        "Wireshark 抓包分析：Display Filter 过滤条件；跟踪 TCP 流查看报文交互；"
        "时间戳与运行日志交叉验证。",
        "argparse 命令行参数解析：替代 sys.argv 位置参数，"
        "使用 --host/--port 等命名参数，更可读、更安全。",
    ]
    for i, k in enumerate(knowledge, 1):
        para(doc, f"{i}. {k}")

    sec(doc, "d.", "git 的 URL")
    para(doc, GIT_URL, indent=False)
    para(doc, "（可推送至 GitHub/Gitee 后替换为远程 URL）", indent=False)

    safe_save(doc, os.path.join(task1_dir, "tcp_packet_capture.docx"))
    print("Task1 done")


def build_task2():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(12)
    s.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    p = doc.add_paragraph()
    p.alignment = 1
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("UDP GBN Reliable Transfer — Wireshark 抓包说明文档")
    set_run(r, "黑体", "Times New Roman", 16)
    r.bold = True

    task2_dir = os.path.join(BASE_DIR, "Task2")

    sec(doc, "a.", "Wireshark 报文捕获截图")
    para(doc, "以下截图展示 UDP 模拟 TCP 可靠传输的完整流程。Display Filter: udp.port == <端口号>")
    cap(doc, "图1 — 三次握手阶段：SYN -> SYN-ACK -> ACK",
        os.path.join(task2_dir, "establish_connection.png"))
    cap(doc, "图2 — 数据传输阶段：DATA 报文 + ACK 累积确认",
        os.path.join(task2_dir, "data_transform.png"))
    cap(doc, "图3 — 丢包与重传：超时重传日志（run_log.txt）",
        os.path.join(task2_dir, "retrans_log.png"))
    cap(doc, "图4 — 四次挥手阶段：FIN -> ACK -> FIN -> ACK",
        os.path.join(task2_dir, "terminate_connection.png"))

    sec(doc, "b.", "实现上的关键点和对应的代码解决方案")
    para(doc, "关键点1：Bit-Flag 协议头设计 + 三次握手。统一 13B 报文头 [1B Flags][4B Seq][4B Ack][4B Length] + Payload。FLAG_SYN=0x01、FLAG_ACK=0x02、FLAG_FIN=0x04，可组合使用。SYN Payload 携带 StudentID（学号后4位 XOR 0x5A3C）进行身份验证。unpack_header() 增加长度检查和 struct.error 兜底，返回 None 统一处理无效报文。代码位于 common.py。", True)
    code(doc, """FLAG_SYN = 0x01; FLAG_ACK = 0x02; FLAG_FIN = 0x04

def pack_header(flags, seq=0, ack=0, length=0):
    return struct.pack("!BIII", flags, seq, ack, length)

def unpack_header(data: bytes):
    if len(data) < 13: return None
    try: return struct.unpack("!BIII", data[:13])
    except struct.error: return None

# 三次握手
Client -> SYN  [Flags=0x01][Seq=0][Ack=0][Len=4][StudentID^MASK, TotalPackets]
Server -> SYN-ACK [Flags=0x03][Seq=0][Ack=1][Len=0]
Client -> ACK  [Flags=0x02][Seq=1][Ack=1][Len=0]

# StudentID = 学号后4位(2913) XOR 0x5A3C = 0x515D
def verify_student_id(received):
    return 0 <= (received ^ config.STUDENT_ID_MASK) <= 9999""")
    para(doc, "关键点2：GBN 滑动窗口 + 累积确认。窗口大小 5 包，base/next_seq 双指针管理。Server 拆分为三个独立阶段方法：_establish_connection() / _receive_data() / _terminate_connection()，由 run() 统一调度，connected 标志位防止握手超时后继续执行。Server 只接受 seq==expected 的包，乱序直接丢弃并重发当前 ACK。累积确认：ACK 值 = 下一期望序号。Client 的 _send_data() 配合独立接收线程 _receiver() 实现发送与 ACK 处理解耦。代码位于 udpserver.py 和 udpclient.py。", True)
    code(doc, """# Server GBN 接收端（udpserver.py ReliableUDPClientHandler）
if seq == expected_seq:
    expected_seq += 1
    send_ack(expected_seq)
else:
    send_ack(expected_seq)       # 乱序丢弃，重发当前 ACK

# Client 滑动窗口（udpclient.py _send_data）
while self.base < config.TOTAL_PACKETS:
    while self.next_seq < self.base + config.WINDOW_SIZE:
        send_packet(self.next_seq)
        self.next_seq += 1
    ack_num = self.ack_queue.get(timeout=self.rto)
    if ack_num > self.base:
        self.base = ack_num""")
    para(doc, "关键点3：EWMA 自适应 RTO + 超时重传 + 快速重传。采用 TCP EWMA 算法：estimated_rtt = 0.875*estimated_rtt + 0.125*sample_rtt，dev_rtt = 0.75*dev_rtt + 0.25*|sample_rtt - estimated_rtt|，RTO = max(50ms, min(3s, (estimated_rtt + 4*dev_rtt)/1000))，初始 RTO=300ms。首次测量直接初始化 estimated_rtt 和 dev_rtt。超时后回退 next_seq 到 base，重传整个窗口。快速重传：独立接收线程持续监听 ACK，连续 3 次相同 ACK 触发 _fast_retransmit()，不等超时立即重传窗口内所有未确认包。Client 废弃全局 settimeout(0.1)，各阶段独立设超时避免干扰 RTO。代码位于 udpclient.py。", True)
    code(doc, """# EWMA 自适应 RTO（udpclient.py _send_data）
sample = self.rtt_samples[-1]
if self.estimated_rtt is None:
    self.estimated_rtt = sample; self.dev_rtt = sample / 2
else:
    self.estimated_rtt = 0.875 * self.estimated_rtt + 0.125 * sample
    self.dev_rtt = 0.75 * self.dev_rtt + 0.25 * abs(sample - self.estimated_rtt)
self.rto = max(0.05, min(3.0, (self.estimated_rtt + 4 * self.dev_rtt) / 1000))

# 快速重传（udpclient.py _receiver）
if ack == self.last_ack:
    self.dup_ack_count += 1
    if self.dup_ack_count == config.FAST_RETX_THRESHOLD:
        self._fast_retransmit()
else:
    self.last_ack = ack; self.dup_ack_count = 0

# 超时重传（udpclient.py _send_data）
except queue.Empty:
    # 超时回退 next_seq，主循环自然重发整个窗口
    self.next_seq = self.base""")
    para(doc, "关键点4：标准四次挥手 + 双线程架构。严格按照 TCP 四次挥手协议：Client 发 FIN → Server 回 ACK → Server 发 FIN → Client 回 ACK。独立接收线程（daemon）使用 queue.Queue 异步传递 ACK 到主线程，实现发送与接收解耦。使用 threading.RLock 可重入锁保护共享状态。挥手前先置 running=False 关闭接收线程，防止抢 socket 的 recvfrom。代码位于 udpclient.py 和 udpserver.py。", True)
    code(doc, """# 双线程架构（udpclient.py ReliableUDPClient）
self.ack_queue = queue.Queue()
self._lock = threading.RLock()
self.recv_thread = threading.Thread(target=self._receiver, daemon=True)

# 各阶段独立超时
_establish_connection():  self.sock.settimeout(0.5)
_receiver():             self.sock.settimeout(0.05)
_terminate_connection(): self.sock.settimeout(0.5)

# 标准四次挥手
Client: 发 FIN [0x04] -> 等 ACK [0x02] -> 等 FIN [0x04] -> 发 ACK [0x02]
Server: 收 FIN -> 发 ACK -> 发 FIN -> 等 ACK""")

    sec(doc, "c.", "掌握的知识点")
    knowledge = [
        "TCP 连接管理：三次握手（SYN->SYN-ACK->ACK）建立连接；"
        "四次挥手（FIN->ACK->FIN->ACK）终止连接的过程及状态转换。",
        "GBN（Go-Back-N）协议：滑动窗口发送机制；累积确认；"
        "超时回退 N 步重传；乱序包直接丢弃而非缓存（与 SR 的区别）。",
        "快速重传（Fast Retransmit）：重复 ACK 检测机制；"
        "3 dup ACK 触发条件；与超时重传的互补关系。",
        "自适应 RTO 计算：TCP EWMA 算法（alpha=0.125, beta=0.25），"
        "estimated_rtt + 4*dev_rtt 公式，上下界裁剪（50ms ~ 3s），"
        "首次 RTT 测量直接初始化避免冷启动偏差。",
        "UDP 编程与可靠性设计：UDP 无连接、不可靠的天然特性；"
        "在应用层通过序列号、ACK、定时器模拟可靠传输；"
        "模拟丢包（random.random() < DROP_RATE）验证协议鲁棒性。",
        "多线程异步 I/O：独立接收线程 + Queue 实现生产者-消费者模式；"
        "daemon 线程生命周期管理；RLock 可重入锁解决嵌套调用死锁；"
        "挥手前停接收线程避免与 FIN 阶段抢 socket recvfrom。",
        "pandas 数据分析：Series.max()/min()/mean()/std() 快速统计 RTT 分布；"
        "丢包率计算公式：(总发送-成功)/总发送x100%；"
        "总发送次数从 packets[i][retrans_count] 聚合计算。",
        "自定义 bit-flag 协议设计：统一 13B 报文头，Flags 位掩码组合，"
        "struct.pack/unpack Big-Endian（!前缀）序列化，"
        "unpack_header 带长度检查和 struct.error 异常兜底。",
    ]
    for i, k in enumerate(knowledge, 1):
        para(doc, f"{i}. {k}")

    sec(doc, "d.", "git 的 URL")
    para(doc, GIT_URL, indent=False)
    para(doc, "（可推送至 GitHub/Gitee 后替换为远程 URL）", indent=False)

    safe_save(doc, os.path.join(task2_dir, "udp_packet_capture.docx"))
    print("Task2 done")


if __name__ == "__main__":
    build_task1()
    build_task2()
    print("All done")
